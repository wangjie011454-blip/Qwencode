import io
from typing import Union


class ResumeParser:
    """简历解析服务"""
    
    def __init__(self):
        self._pdf_available = False
        self._docx_available = False
        
        # 尝试导入 PDF 解析库
        try:
            import pdfplumber
            self._pdfplumber = pdfplumber
            self._pdf_available = True
        except ImportError:
            pass
        
        # 尝试导入 Word 解析库
        try:
            import docx
            self._docx = docx
            self._docx_available = True
        except ImportError:
            pass
    
    def parse(self, content: bytes, file_type: str) -> str:
        """
        解析简历文件
        
        Args:
            content: 文件二进制内容
            file_type: 文件类型 (pdf, docx, doc)
            
        Returns:
            解析后的纯文本
        """
        if file_type == "pdf":
            return self._parse_pdf(content)
        elif file_type in ["docx", "doc"]:
            return self._parse_docx(content)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")
    
    def _parse_pdf(self, content: bytes) -> str:
        """解析 PDF 文件"""
        if not self._pdf_available:
            raise ImportError("pdfplumber 库未安装，请运行: pip install pdfplumber")

        text_parts = []
        with self._pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

        return "\n".join(text_parts)
    
    def _parse_docx(self, content: bytes) -> str:
        """解析 Word 文件"""
        if not self._docx_available:
            raise ImportError("python-docx 库未安装，请运行: pip install python-docx")
        
        doc = self._docx.Document(io.BytesIO(content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 尝试提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)